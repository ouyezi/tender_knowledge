from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches

from src.config import Settings
from src.models.actual_bid_parse_task import ActualBidParseTask, ActualBidParseTaskStatus
from src.models.candidate_knowledge import CandidateKnowledge
from src.models.chapter_taxonomy import CategoryStatus, ChapterTaxonomy
from src.models.document_media_asset import DocumentMediaAsset
from src.models.document_tree_node import DocumentTreeNode, DocumentTreeNodeType
from src.models.downstream_task_entry import (
    DownstreamTaskEntry,
    DownstreamTaskStatus,
    DownstreamTaskType,
)
from src.models.file_import import FileImport, FileImportStatus, FilePurpose, FileType
from src.services.actual_bid_parse_runner import run_actual_bid_parse_pending
from src.services.candidate_generate_service import generate_for_document
from src.services.content_blocks import parse_content

_MINIMAL_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
    "0000000a49444154789c63000100000500010d0a2ddb0000000049454e44ae426082"
)


def _make_docx_with_heading_and_image(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("技术方案", level=1)
    doc.add_paragraph("章节正文")
    doc.add_paragraph().add_run().add_picture(BytesIO(_MINIMAL_PNG), width=Inches(1.0))
    doc.save(path)


def _seed_confirmed_actual_bid_import_with_downstreams(db_session, seeded_kb, source_docx: Path):
    storage_rel = f"{seeded_kb.kb_id}/sample-visibility.docx"
    storage_abs = Path(Settings().storage_root) / storage_rel
    storage_abs.parent.mkdir(parents=True, exist_ok=True)
    storage_abs.write_bytes(source_docx.read_bytes())

    record = FileImport(
        kb_id=seeded_kb.kb_id,
        file_name="sample-visibility.docx",
        file_type=FileType.docx,
        file_size=storage_abs.stat().st_size,
        storage_path=storage_rel,
        status=FileImportStatus.confirmed,
        file_purpose=FilePurpose.actual_bid,
        created_by="admin",
    )
    db_session.add(record)
    db_session.flush()
    for task_type in (
        DownstreamTaskType.document_parse,
        DownstreamTaskType.bid_outline_extract,
        DownstreamTaskType.candidate_knowledge_generate,
    ):
        db_session.add(
            DownstreamTaskEntry(
                kb_id=seeded_kb.kb_id,
                import_id=record.import_id,
                task_type=task_type,
                status=DownstreamTaskStatus.pending,
            )
        )
    db_session.commit()
    return record


def test_parse_docx_image_links_tree_and_candidate_blocks(db_session, seeded_kb, tmp_path):
    source_docx = tmp_path / "with-image.docx"
    _make_docx_with_heading_and_image(source_docx)
    record = _seed_confirmed_actual_bid_import_with_downstreams(db_session, seeded_kb, source_docx)

    run_actual_bid_parse_pending(db_session)

    task = (
        db_session.query(ActualBidParseTask)
        .filter(ActualBidParseTask.import_id == record.import_id)
        .order_by(ActualBidParseTask.created_at.desc())
        .first()
    )
    assert task is not None
    assert task.status == ActualBidParseTaskStatus.ready
    assert task.document_id is not None

    image_nodes = (
        db_session.query(DocumentTreeNode)
        .filter(
            DocumentTreeNode.document_id == task.document_id,
            DocumentTreeNode.node_type == DocumentTreeNodeType.image,
        )
        .all()
    )
    assert image_nodes
    assert all(node.content_ref for node in image_nodes)
    image_asset_ids = {str(node.content_ref) for node in image_nodes if node.content_ref}

    media_assets = (
        db_session.query(DocumentMediaAsset)
        .filter(DocumentMediaAsset.document_id == task.document_id)
        .all()
    )
    assert media_assets
    assert {str(asset.asset_id) for asset in media_assets}.issuperset(image_asset_ids)
    assert all((Path(Settings().storage_root) / asset.storage_path).is_file() for asset in media_assets)

    taxonomy = ChapterTaxonomy(
        kb_id=seeded_kb.kb_id,
        parent_id=None,
        standard_name="技术方案",
        taxonomy_code="technical_solution",
        status=CategoryStatus.active,
        path="",
        depth=0,
    )
    db_session.add(taxonomy)
    db_session.flush()
    taxonomy.path = f"/{taxonomy.taxonomy_id}/"

    heading = (
        db_session.query(DocumentTreeNode)
        .filter(
            DocumentTreeNode.document_id == task.document_id,
            DocumentTreeNode.node_type == DocumentTreeNodeType.heading,
        )
        .order_by(DocumentTreeNode.sort_order.asc())
        .first()
    )
    assert heading is not None
    heading.chapter_taxonomy_id = taxonomy.taxonomy_id
    db_session.query(CandidateKnowledge).filter(
        CandidateKnowledge.import_id == record.import_id,
        CandidateKnowledge.source_doc_id == task.document_id,
    ).delete(synchronize_session=False)
    db_session.commit()

    created = generate_for_document(
        db_session,
        kb_id=seeded_kb.kb_id,
        import_id=record.import_id,
        document_id=task.document_id,
        parse_task_id=task.parse_task_id,
    )
    assert created
    parsed = parse_content(created[0].content)
    assert any(
        block.get("type") == "image" and str(block.get("asset_id") or "") in image_asset_ids
        for block in parsed.blocks
    )
