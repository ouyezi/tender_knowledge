from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Inches

from src.services.docx_image_extractor import extract_docx_images

# Minimal 1x1 PNG (no Pillow dependency)
_MINIMAL_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
    "0000000a49444154789c63000100000500010d0a2ddb0000000049454e44ae426082"
)


def _make_docx_with_image(path: Path) -> None:
    buf = BytesIO(_MINIMAL_PNG)
    doc = Document()
    doc.add_paragraph("before")
    doc.add_paragraph().add_run().add_picture(buf, width=Inches(1.0))
    doc.save(path)


def test_extract_docx_images_writes_storage_and_records(tmp_path):
    docx_path = tmp_path / "with-image.docx"
    _make_docx_with_image(docx_path)
    kb_id = uuid4()
    document_id = uuid4()
    storage_root = tmp_path / "storage"

    results = extract_docx_images(
        docx_path,
        storage_root=storage_root,
        kb_id=kb_id,
        document_id=document_id,
    )
    assert len(results) >= 1
    assert results[0].asset_id is not None
    assert (storage_root / results[0].storage_path).is_file()
    assert results[0].mime_type.startswith("image/")
